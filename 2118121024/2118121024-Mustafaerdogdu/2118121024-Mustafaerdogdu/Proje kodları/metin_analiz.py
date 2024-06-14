import re
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from docx import Document
import nltk
import pickle
from snowballstemmer import stemmer
from gensim.models import KeyedVectors
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score
import numpy as np
from sklearn.model_selection import train_test_split
from gensim.models import Word2Vec


# NLTK'nin gerekli veri kaynaklarını indirin
nltk.download('punkt')
nltk.download('stopwords')

# İngilizce stopwords listesini tanımlayın
en_stopwords = set(stopwords.words('english'))

# Word belgesinden metin verilerini okuyan fonksiyon
def word_veri_oku(dosya_yolu):
    doc = Document(dosya_yolu)
    metinler = []
    for paragraf in doc.paragraphs:
        metinler.append(paragraf.text)
    return metinler

# Metni temizleme fonksiyonunu tanımlayın
def metni_temizle(metin, stopwordler):
    metin = re.sub(r'\$\w*', '', metin)  # Para birimlerini temizle
    metin = re.sub(r'^RT[\s]+', '', metin)  # Retweet işaretlerini temizle
    metin = re.sub(r'https?:\/\/.*[\r\n]*', '', metin)  # URL'leri temizle
    metin = re.sub(r'#', '', metin)  # Hashtag işaretlerini temizle
    metin = re.sub(r':\(\)', '', metin)  # Emojileri temizle
    metin = re.sub(r'[.,!?;:•]', '', metin)  # Noktalama işaretlerini temizle
    
    metin = metin.lower()
    
    # Cümleleri kelimelere ayır
    tokenizer = word_tokenize(metin)
    
    # Temizlenmiş kelimeleri saklayacak liste
    temizlenmis_kelimeler = []
    
    # Her kelimeyi kontrol et
    for kelime in tokenizer:
        if kelime.lower() not in stopwordler:
            kelime_koku = stemmer('english').stemWord(kelime)
            temizlenmis_kelimeler.append(kelime_koku)
            
    return temizlenmis_kelimeler

# Word belgesinin yolu
dosya_yolu = "kelimeler.docx"

# Word belgesinden metin verilerini oku
kendi_metinler = word_veri_oku(dosya_yolu)

# Metinleri temizle
temizlenmis_kelimeler = []
for metin in kendi_metinler:
    temiz_metin = metni_temizle(metin, en_stopwords)
    temizlenmis_kelimeler.extend(temiz_metin)

# Kullanıcıdan cümle al
girilen_cumle = input("Lütfen bir cümle girin: ")

# Temizleme işlemini gerçekleştir
temizlenmis_cümle = metni_temizle(girilen_cumle, en_stopwords)

# Temizlenmiş kelimeleri yazdır
#print("Temizlenmiş kelimeler:", temizlenmis_kelimeler)

# Temizlenmiş cümleyi ekrana yazdır
#print("Temizlenmiş cümle:", temizlenmis_cümle)

def kelime_sayisi(word_dosya_yolu, kelimeler):
    doc = Document(word_dosya_yolu)
    kelime_sayac = {}
    for kelime in kelimeler:
        kelime_sayac[kelime] = 0
    for paragraf in doc.paragraphs:
        for kelime in kelimeler:
            kelime_sayac[kelime] += paragraf.text.lower().count(kelime.lower())
    return kelime_sayac

# Word belgesinin yolu
word_dosya_yolu = "kelimeler.docx"

# Her kelimenin kelime sayısını bul ve ekrana yazdır
kelime_sayac = kelime_sayisi(word_dosya_yolu, temizlenmis_cümle)

# Filtrelenmiş kelimeleri ve belgedeki kullanım sayılarını ekrana yazdır
for kelime in temizlenmis_cümle:
    sayi = kelime_sayac.get(kelime, 0)
    print(f"'{kelime}' kelimesi Word belgesinde {sayi} kez kullanılmıştır.")

# Gruplara ayırmak için kelimeleri listelerde saklayın
mountain = ["mountain"]
drop = ["drop"]
sea = ["sea"]
snow = ["snow"]
fog = ["fog"]
green_area = ["green", "area"]

# Grupları saymak için sayaçlar oluşturun
sayac_mountain = 0
sayac_drop = 0
sayac_sea = 0
sayac_snow = 0
sayac_fog = 0
sayac_green_area = 0

# Temizlenmiş cümleyi küçük harflere dönüştürün
temizlenmis_cümle_lower = girilen_cumle.lower()

# Her kelimeyi kontrol edin ve uygun gruba ekleyin
for kelime in temizlenmis_cümle_lower.split():
    if kelime in mountain:
        sayac_mountain += 1
    elif kelime in drop:
        sayac_drop += 1
    elif kelime in sea:
        sayac_sea += 1
    elif kelime in snow:
        sayac_snow += 1
    elif kelime in fog:
        sayac_fog += 1
    elif kelime in green_area:
        sayac_green_area += 1

# Gruplara ayrılan kelimeleri ekrana yazdırın
print("Mountain:", sayac_mountain)
print("Drop:", sayac_drop)
print("Sea:", sayac_sea)
print("Snow:", sayac_snow)
print("Fog:", sayac_fog)
print("Green Area:", sayac_green_area)



# Önceden eğitilmiş kelime gömme modelini yükleme
word_vectors = KeyedVectors.load_word2vec_format('GoogleNews-vectors-negative300.bin', binary=True)

# Metin verilerini kelime gömme vektörlerine dönüştürme
def word_veri_oku(dosya_yolu):
    doc = Document(dosya_yolu)
    metinler = []
    for paragraf in doc.paragraphs:
        if paragraf.text.strip():  # Boş paragrafları atla
            metinler.append(paragraf.text)
    return metinler


# Word belgesinin yolu
dosya_yolu = "kelimeler.docx"

# Word belgesinden metin verilerini oku
texts = word_veri_oku(dosya_yolu)


labels = [1] * 120 + [0] * 21


print(f"Toplam metin sayısı: {len(texts)}")
print(f"Pozitif metin sayısı: {len([label for label in labels if label == 1])}")
print(f"Negatif metin sayısı: {len([label for label in labels if label == 0])}")


# Veriyi eğitim ve test setlerine ayırma
X_train, X_test, y_train, y_test = train_test_split(texts, labels, test_size=0.2, random_state=42)

# Metin verilerini kelime gömme vektörlerine dönüştürme fonksiyonu
def text_to_vector(text):
    words = text.split()
    vectors = []
    for word in words:
        if word in word_vectors:
            vectors.append(word_vectors[word])
    return sum(vectors) / len(vectors) if vectors else np.zeros(300)

# Eğitim ve test verilerini kelime gömme vektörlerine dönüştürme
X_train_vec = np.array([text_to_vector(text) for text in X_train])
X_test_vec = np.array([text_to_vector(text) for text in X_test])

# Modeli eğitme
model = LogisticRegression()
model.fit(X_train_vec, y_train)

# Modeli değerlendirme
y_pred = model.predict(X_test_vec)
accuracy = accuracy_score(y_test, y_pred)
print("Accuracy:", accuracy)

# Modeli kaydetme
with open('model.pkl', 'wb') as f:
    pickle.dump(model, f)

# Modeli yükleme
with open('model.pkl', 'rb') as f:
    loaded_model = pickle.load(f)
 

# Kullanıcıdan iki kelime girişi al
kelime1 = input("Lütfen bir kelime girin: ")
kelime2 = input("Lütfen başka bir kelime girin: ")

# Kelimelerin benzerliğini hesapla
benzerlik = word_vectors.similarity(kelime1, kelime2)
print(f"{kelime1} ve {kelime2} kelimelerinin benzerliği: {benzerlik}")


def text_to_vectors(text, vector_size=150, window=10, min_count=1, sg=1):
    # Metni kelime düzeyinde tokenize etme
    tokens = word_tokenize(text.lower())  # Küçük harfe dönüştürme önemli olabilir
    
    # Word2Vec modelini yaratma ve eğitme
    model = Word2Vec([tokens], vector_size=vector_size, window=window, min_count=min_count, sg=sg)
    
    # Kelimelerin vektörlerini alma
    word_vectors = {}
    for word in tokens:
        word_vectors[word] = model.wv[word]
    
    return word_vectors

# Kullanıcıdan metni al
user_text = input("Lütfen kelime girin: ")

# Metni vektörlere dönüştürme
vectors = text_to_vectors(user_text)

# Her kelimenin vektörünü gösterme
for word, vector in vectors.items():
    print(f"{word}: {vector}")




